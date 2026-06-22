"""
tests/test_office.py
Tests unitarios para el motor Office y el motor de texto.
"""

import sys
import unittest
import tempfile
import shutil
import json
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.compression_profile import QualityLevel
from compressors.base import CompressionOptions


class TestOfficeCompressor(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.tmp_path = Path(self.tmpdir)

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _create_test_docx(self) -> Path:
        """Crea un DOCX de prueba con python-docx."""
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Documento de Prueba", 0)
            for i in range(20):
                doc.add_paragraph(f"Este es el párrafo número {i + 1} del documento de prueba.")
            path = self.tmp_path / "test.docx"
            doc.save(str(path))
            return path
        except ImportError:
            self.skipTest("python-docx no instalado")

    def test_docx_compression(self):
        try:
            from compressors.office.office_compressor import OfficeCompressor
            c = OfficeCompressor()
            docx_path = self._create_test_docx()
            output = self.tmp_path / "out.docx"
            opts = CompressionOptions(
                quality_level=QualityLevel.BALANCED,
                output_path=output,
                remove_metadata=True,
            )
            result = c.compress(docx_path, opts)
            self.assertIsNotNone(result)
            self.assertGreater(result.original_size, 0)
        except ImportError as e:
            self.skipTest(str(e))

    def test_metadata_cleaning(self):
        """La función de limpieza XML no debe lanzar excepciones."""
        from compressors.office.office_compressor import _clean_metadata_xml

        sample_xml = b"""<?xml version="1.0" encoding="UTF-8"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties">
    <dc:creator xmlns:dc="http://purl.org/dc/elements/1.1/">Juan Perez</dc:creator>
    <cp:lastModifiedBy>Maria Garcia</cp:lastModifiedBy>
</cp:coreProperties>"""

        result = _clean_metadata_xml(sample_xml)
        self.assertIsInstance(result, bytes)
        self.assertGreater(len(result), 0)


class TestTextCompressor(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.tmp_path = Path(self.tmpdir)

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_json_minification(self):
        from compressors.text.text_compressor import TextCompressor
        c = TextCompressor()

        data = {"nombre": "prueba", "valor": 42, "lista": [1, 2, 3]}
        json_path = self.tmp_path / "data.json"
        with open(json_path, "w") as f:
            json.dump(data, f, indent=4)  # Indentado (grande)

        output = self.tmp_path / "data_optimizado.json"
        opts = CompressionOptions(quality_level=QualityLevel.BALANCED, output_path=output)
        result = c.compress(json_path, opts)

        self.assertTrue(result.success)
        if output.exists():
            # El minificado debe ser más pequeño
            minified = output.read_text()
            self.assertNotIn("    ", minified)  # Sin indentación

    def test_html_minification(self):
        from compressors.text.text_compressor import TextCompressor
        c = TextCompressor()

        html_content = """<!DOCTYPE html>
<html>
    <head>
        <!-- Este es un comentario -->
        <title>Prueba</title>
    </head>
    <body>
        <h1>Hola Mundo</h1>
        <p>Este es un párrafo.</p>
    </body>
</html>"""

        html_path = self.tmp_path / "test.html"
        html_path.write_text(html_content, encoding="utf-8")

        output = self.tmp_path / "test_optimizado.html"
        opts = CompressionOptions(quality_level=QualityLevel.BALANCED, output_path=output)
        result = c.compress(html_path, opts)

        self.assertTrue(result.success)
        if output.exists():
            minified = output.read_text()
            self.assertNotIn("<!-- Este es un comentario -->", minified)

    def test_text_cleaning(self):
        from compressors.text.text_compressor import TextCompressor
        c = TextCompressor()

        txt = "Línea uno   \n\n\n\nLínea dos  \n\n\n\nLínea tres\n"
        txt_path = self.tmp_path / "test.txt"
        txt_path.write_text(txt, encoding="utf-8")

        output = self.tmp_path / "test_optimizado.txt"
        opts = CompressionOptions(quality_level=QualityLevel.BALANCED, output_path=output)
        result = c.compress(txt_path, opts)
        self.assertTrue(result.success)


class TestHistoryManager(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.db_path = Path(self.tmpdir) / "test_history.db"

    def tearDown(self):
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_add_and_retrieve(self):
        from utils.history_manager import HistoryManager
        from models.history_entry import HistoryEntry

        mgr = HistoryManager(db_path=self.db_path)
        entry = HistoryEntry(
            filename="test.pdf",
            file_type="pdf",
            original_size=1_000_000,
            final_size=400_000,
            reduction_percent=60.0,
            elapsed_seconds=1.5,
            output_path="/tmp/test_optimizado.pdf",
            method_used="Test",
            success=True,
        )
        mgr.add_entry(entry)
        entries = mgr.get_all()
        self.assertEqual(len(entries), 1)
        self.assertEqual(entries[0].filename, "test.pdf")
        self.assertAlmostEqual(entries[0].reduction_percent, 60.0)

    def test_clear(self):
        from utils.history_manager import HistoryManager
        from models.history_entry import HistoryEntry

        mgr = HistoryManager(db_path=self.db_path)
        entry = HistoryEntry(
            filename="a.pdf", file_type="pdf",
            original_size=1000, final_size=500,
            reduction_percent=50.0, elapsed_seconds=0.5,
            output_path="", method_used="Test", success=True,
        )
        mgr.add_entry(entry)
        mgr.clear()
        self.assertEqual(len(mgr.get_all()), 0)


if __name__ == "__main__":
    unittest.main()
